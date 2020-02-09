import re
import docx
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor



f = open ('input_goto2.txt')
lines = f.read().splitlines()

str = f.read()
docu = Document()




start = "1"
last_goto = "1"
for line in lines:
    

    run = docu.add_paragraph().add_run(line)
    
            
        
    for word in line.split():

        if(word == last_goto):
            start = "1"
            break
        
        if (word == 'goto' and start == "1"):
            run.font.color.rgb = RGBColor(255, 0, 0)
            last_goto = line.split("goto",1)[1]
            last_goto = last_goto.split()[0]
            start = "0"      
            break

#docu.save('output.docx')




def voidLabel(strr):

    match = re.search(r'void (\S+)',strr)
    if match:
        print(match)
        namevar = match.group(1)
        for i in namevar:
            if (i=='('):
                break
            newstr1+=i

        return newstr1

voidLabel(str)

'''
list = ["void fun()", "void funt()", "gala se"]
for e in list:
    z = re.match("void (\S+)",e)

if z:
    print(z.group(1))

'''

emails = re.findall(r'[void] [\S+]', str)
for email in emails:
    print(email)


'''
match = re.search(r'(double|int|char|float) (\S+)',str)
print(match)
namevar = match.group(2)
for i in namevar:
    if (i=='('):
        break
    newstr2+=i   

'''
     
     




#print(str)

#void (\S+)


#(double|int|char|float) (\S+)






def getJumplabelHighLight(str):
   match = re.search(r'goto (\S+);',str)
   if match:
       print(match)
       label = match.group(1)
       label=label+':'
       # print( label)
       return  label

def getallItemsinLabel(x,lab):
    match = re.search(r'(?<='+lab+')[^}]*',x)
    content= match.group(0)
    return  content
    #print(content)


def getAllprints(xx):
    match = re.findall(r'printf',xx)
    l = len(match)
    print("number of printf are",l)

getJumplabelHighLight(str)

'''
label=getJumplabel(str)
getallItemsinLabel(str,label)
getAllprints(str)
getAllprints(getallItemsinLabel(str,getJumplabel(str)))
print(getJumplabel(str))
'''
f.close()